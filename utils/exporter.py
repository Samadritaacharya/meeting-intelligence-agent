"""Export meeting analysis results to PDF, DOCX, CSV, and Markdown."""
from __future__ import annotations

import csv
from io import BytesIO, StringIO
from typing import Any, Dict, Iterable, List


def _as_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _stringify(value: Any) -> str:
    if isinstance(value, dict):
        return " | ".join(f"{k}: {v}" for k, v in value.items())
    return str(value)


def _action_rows(result: Dict[str, Any]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for item in _as_list(result.get("action_items")):
        if isinstance(item, dict):
            rows.append(
                {
                    "owner": str(item.get("owner") or "TBD"),
                    "action": str(item.get("action") or item.get("description") or item.get("summary") or "Review follow-up item"),
                    "due_date": str(item.get("due_date") or item.get("deadline") or "TBD"),
                    "priority": str(item.get("priority") or "Medium"),
                    "status": str(item.get("status") or "Open"),
                }
            )
        else:
            rows.append({"owner": "TBD", "action": str(item), "due_date": "TBD", "priority": "Medium", "status": "Open"})
    return rows


def _risk_rows(result: Dict[str, Any]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for item in _as_list(result.get("risks_flagged")):
        if isinstance(item, dict):
            rows.append(
                {
                    "risk": str(item.get("risk") or item.get("description") or "Unspecified risk"),
                    "impact": str(item.get("impact") or "Medium"),
                    "likelihood": str(item.get("likelihood") or "Medium"),
                    "mitigation": str(item.get("mitigation") or "Define mitigation and owner."),
                    "owner": str(item.get("owner") or "TBD"),
                }
            )
        else:
            rows.append(
                {
                    "risk": str(item),
                    "impact": "Medium",
                    "likelihood": "Medium",
                    "mitigation": "Define mitigation and owner.",
                    "owner": "TBD",
                }
            )
    return rows


def export_to_pdf(result: Dict[str, Any]) -> bytes:
    """Return a meeting intelligence report as PDF bytes."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 18)
    pdf.cell(0, 12, "Meeting Intelligence Report", ln=True)
    pdf.ln(3)

    def section(title: str, content: Any) -> None:
        pdf.set_font("Arial", "B", 13)
        pdf.cell(0, 9, title.encode("latin-1", "replace").decode("latin-1"), ln=True)
        pdf.set_font("Arial", "", 10)
        if isinstance(content, list):
            for item in content:
                pdf.multi_cell(0, 6, ("- " + _stringify(item)).encode("latin-1", "replace").decode("latin-1"))
        elif isinstance(content, dict):
            for k, v in content.items():
                pdf.multi_cell(0, 6, f"{k}: {v}".encode("latin-1", "replace").decode("latin-1"))
        else:
            pdf.multi_cell(0, 6, str(content).encode("latin-1", "replace").decode("latin-1"))
        pdf.ln(2)

    section("Meeting Type", result.get("meeting_type", ""))
    section("Executive Summary", result.get("executive_summary", ""))
    section("PMO Status", result.get("pmo_status", {}))
    section("Key Decisions", result.get("key_decisions", []))
    section("Action Items", _action_rows(result))
    section("Risks Flagged", _risk_rows(result))
    section("Open Questions", result.get("open_questions", []))
    section("Follow-up Email", result.get("follow_up_email", {}))
    data = pdf.output(dest="S")
    return data if isinstance(data, bytes) else data.encode("latin-1")


def export_to_docx(result: Dict[str, Any]) -> bytes:
    """Return a meeting intelligence report as DOCX bytes."""
    from docx import Document

    doc = Document()
    doc.add_heading("Meeting Intelligence Report", level=0)

    def section(title: str, content: Any) -> None:
        doc.add_heading(title, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(_stringify(item), style="List Bullet")
        elif isinstance(content, dict):
            for k, v in content.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(content))

    section("Meeting Type", result.get("meeting_type", ""))
    section("Executive Summary", result.get("executive_summary", ""))
    section("PMO Status", result.get("pmo_status", {}))
    section("Key Decisions", result.get("key_decisions", []))
    section("Action Items", _action_rows(result))
    section("Risks Flagged", _risk_rows(result))
    section("Open Questions", result.get("open_questions", []))
    section("Follow-up Email", result.get("follow_up_email", {}))
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_actions_to_jira_csv(result: Dict[str, Any]) -> bytes:
    """Export action items in a Jira-import-friendly CSV format."""
    output = StringIO()
    fieldnames = ["Summary", "Description", "Issue Type", "Priority", "Assignee", "Due Date", "Status"]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for row in _action_rows(result):
        writer.writerow(
            {
                "Summary": row["action"][:90],
                "Description": row["action"],
                "Issue Type": "Task",
                "Priority": row["priority"],
                "Assignee": row["owner"],
                "Due Date": row["due_date"],
                "Status": row["status"],
            }
        )
    return output.getvalue().encode("utf-8")


def export_risks_to_csv(result: Dict[str, Any]) -> bytes:
    """Export extracted risks to CSV."""
    output = StringIO()
    fieldnames = ["Risk", "Impact", "Likelihood", "Mitigation", "Owner"]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for row in _risk_rows(result):
        writer.writerow(
            {
                "Risk": row["risk"],
                "Impact": row["impact"],
                "Likelihood": row["likelihood"],
                "Mitigation": row["mitigation"],
                "Owner": row["owner"],
            }
        )
    return output.getvalue().encode("utf-8")


def export_to_markdown(result: Dict[str, Any]) -> bytes:
    """Export a Notion/GitHub-friendly Markdown meeting note."""
    lines: List[str] = ["# Meeting Intelligence Report", ""]
    lines.append(f"**Meeting type:** {result.get('meeting_type', 'Project sync')}")
    status = result.get("pmo_status", {}) or {}
    if status:
        lines.extend(
            [
                "",
                "## PMO Status",
                f"- **RAG status:** {status.get('rag_status', 'TBD')}",
                f"- **Overall health:** {status.get('overall_health', 'TBD')}",
                f"- **Decision needed:** {status.get('decision_needed', 'TBD')}",
                f"- **Next review:** {status.get('next_review', 'TBD')}",
            ]
        )

    lines.extend(["", "## Executive Summary", str(result.get("executive_summary", "")), "", "## Key Decisions"])
    for item in _as_list(result.get("key_decisions")):
        lines.append(f"- {_stringify(item)}")

    lines.extend(["", "## Action Items"])
    for row in _action_rows(result):
        lines.append(f"- **{row['owner']}** — {row['action']} — Due: {row['due_date']} — Priority: {row['priority']} — Status: {row['status']}")

    lines.extend(["", "## Risks"])
    for row in _risk_rows(result):
        lines.append(f"- **{row['impact']} impact / {row['likelihood']} likelihood:** {row['risk']} — Mitigation: {row['mitigation']} — Owner: {row['owner']}")

    lines.extend(["", "## Open Questions"])
    for item in _as_list(result.get("open_questions")):
        lines.append(f"- {_stringify(item)}")

    email = result.get("follow_up_email", {}) or {}
    lines.extend(["", "## Follow-up Email", f"**Subject:** {email.get('subject', '')}", "", str(email.get("body", "")).replace("<br>", "\n")])
    return "\n".join(lines).encode("utf-8")
